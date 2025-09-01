# -*- coding: utf-8 -*-
from __future__ import annotations
from typing import List
import pandas as pd
from docx import Document
from docx.text.paragraph import Paragraph
from models import Item

def export_stock_to_excel(items: List[Item], path: str) -> None:
    rows = []
    for it in items:
        rows.append({
            "SEQ_ID": it.seq_id,
            "Наименование": it.name,
            "Категория": it.category,
            "Количество": it.quantity,
            "Ед. изм.": it.unit,
            "Фасовка": it.packaging,
            "Место хранения": it.storage_place,
            "Срок годности": it.expiry_date or "",
            "Дата поступления": it.date_received or "",
            "Номер партии": it.batch_number,
            "Ответственный": it.responsible,
            "Квалификация": it.qualification or "",
            "Тип реактива": it.reagent_type or "",
            "№ в госреестре СО": it.state_register_no or "",
            "Аттестованное значение": it.certified_value or "",
            "Дата выпуска": it.manufacture_date or "",
            "Производитель": it.manufacturer or "",
            "Условия хранения": it.storage_conditions or ""
        })
    pd.DataFrame(rows).to_excel(path, index=False)

def _replace_in_paragraph(paragraph: Paragraph, mapping: dict):
    inline = paragraph.runs
    if not inline:
        return
    text = "".join(run.text for run in inline)
    for k, v in mapping.items():
        if k in text:
            text = text.replace(k, str(v))
    paragraph.text = text

def _replace_in_tables(doc: Document, mapping: dict):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)

def export_issue_docx(item: Item, path: str, template_path: str) -> None:
    mapping = {
        "{ID}": item.seq_id,
        "{SEQ_ID}": item.seq_id,
        "{NAME}": item.name,
        "{CATEGORY}": item.category,
        "{QTY}": item.quantity,
        "{UNIT}": item.unit,
        "{BATCH}": item.batch_number,
        "{RESPONSIBLE}": item.responsible,
    }
    doc = Document(template_path)
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)
    _replace_in_tables(doc, mapping)
    doc.save(path)